"""
OFA Policy Brief & Evidence Dossier Generator Module
Synthesizes project dossiers, empirical parameters, and literature into publication-ready policy briefs,
board memos, and technical appendices with Word (.docx), Markdown (.md), and HTML exports.
"""
import os
import io
import re
from db import (
    get_connection, get_project, get_project_papers,
    get_empirical_parameters, get_project_syntheses
)


BRIEF_TEMPLATES = {
    "Government Submission / Policy Brief": {
        "audience": "Ontario Ministry of Agriculture, Food and Agribusiness (OMAFA), AAFC, Standing Committees",
        "description": "Formal policy brief with strategic problem statement, empirical literature consensus, farm financial impact, and actionable recommendations."
    },
    "Executive Board Memo / Decision Brief": {
        "audience": "OFA Board of Directors, Executive Committee, Policy Advisory Council (PAC)",
        "description": "High-level strategic memo emphasizing farm-gate risks, commodity exposures, and recommended OFA policy stances."
    },
    "Technical Economic Appendix & Parameter Dossier": {
        "audience": "Agricultural Economists, Policy Analysts, Technical Submissions",
        "description": "Rigorous technical appendix detailing econometric methodologies, elasticity distributions, IO multipliers, and simulator assumptions."
    }
}


def assemble_dossier_context(conn, project_id: int) -> dict:
    """Compile all literature, empirical parameters, and syntheses for a project."""
    project = get_project(conn, project_id)
    if not project:
        raise ValueError(f"Project ID {project_id} not found.")

    papers = get_project_papers(conn, project_id)
    syntheses = get_project_syntheses(conn, project_id)
    params = get_empirical_parameters(conn, project_id=project_id)

    # If project-specific params are few, include relevant baseline params
    if len(params) < 3:
        all_p = get_empirical_parameters(conn)
        params = params + all_p[:12]

    return {
        "project": project,
        "papers": papers,
        "syntheses": syntheses,
        "parameters": params
    }


def generate_policy_brief(
    project_id: int,
    template_type: str,
    target_audience: str = "",
    custom_focus: str = "",
    api_key: str = None
) -> str:
    """Generate a structured policy brief or board memo using Gemini."""
    if not api_key:
        api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key:
        raise ValueError("A Gemini API Key is required to generate policy briefs.")

    import google.generativeai as genai
    genai.configure(api_key=api_key)

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
    except Exception:
        model = genai.GenerativeModel("gemini-1.5-flash")

    conn = get_connection()
    context_data = assemble_dossier_context(conn, project_id)
    conn.close()

    proj = context_data["project"]
    papers = context_data["papers"]
    params = context_data["parameters"]
    syntheses = context_data["syntheses"]

    # 1. Format Paper Summaries
    paper_entries = []
    for i, p in enumerate(papers[:30], 1):
        core_tag = "[CORE REFERENCE] " if p.get("is_core") else ""
        tag_str = f"Tag: {p.get('relevance_tag')} | " if p.get("relevance_tag") else ""
        note_str = f"Analyst Note: {p.get('analyst_notes')} | " if p.get("analyst_notes") else ""
        src = p.get("source_name") or "Academic Source"
        abstract = (p.get("abstract") or "")[:400]
        entry = f"""[{i}] {core_tag}"{p.get('title')}" ({p.get('year', 'N/A')})
    Source: {src} | Citations: {p.get('citation_count', 0)} | {tag_str}{note_str}DOI: {p.get('doi', 'N/A')}
    Abstract: {abstract}"""
        paper_entries.append(entry)

    papers_text = "\n---\n".join(paper_entries)

    # 2. Format Parameters
    param_entries = []
    for p in params:
        range_s = f" [Range: {p['stat_lower']} to {p['stat_upper']}]" if p.get("stat_lower") is not None else ""
        se_s = f" (SE: ±{p['standard_error']})" if p.get("standard_error") is not None else ""
        param_entries.append(
            f"* **{p['commodity']} - {p['parameter_type']}**: {p['point_estimate']} {p.get('unit','')}{range_s}{se_s} (Model: {p.get('model_type','Econometric')}, Horizon: {p.get('time_horizon','N/A')}) - *{p.get('notes','')}*"
        )
    params_text = "\n".join(param_entries) if param_entries else "No empirical parameters registered."

    # 3. Prior Synthesis Text
    prior_synth_text = syntheses[0]["synthesis_markdown"][:2000] if syntheses else "None."

    # 4. Template Instructions
    if template_type == "Executive Board Memo / Decision Brief":
        template_instructions = """
STRUCTURE AS AN OFA EXECUTIVE BOARD DECISION BRIEF:
1. **Header Block**: TO: OFA Board of Directors & Policy Advisory Council | FROM: Economic & Policy Analysis Division | DATE: [Current Date] | SUBJECT: [Clear Title]
2. **Executive Summary & Strategic Context**: The core issue, urgency, and why this matters to Ontario farm businesses.
3. **Key Economic & Farm-Level Findings**: Synthesis of empirical evidence, commodity-specific exposures, and cash flow risk magnitudes.
4. **Calibrated Impact Summary**: Key numbers from our economic models (pass-through rates, revenue multipliers, cost burdens).
5. **Strategic Policy Risks & Recommended OFA Stance**: Defensible, actionable resolutions and advocacy positions for the Board to adopt.
"""
    elif template_type == "Technical Economic Appendix & Parameter Dossier":
        template_instructions = """
STRUCTURE AS A TECHNICAL ECONOMIC APPENDIX & PARAMETER DOSSIER:
1. **Title & Methodological Scope**: Analytical framework and econometric foundations.
2. **Econometric Parameter Distribution Matrix**: Markdown table comparing point estimates, confidence intervals, sample periods, and estimation models across studies.
3. **Simulation Calibration Assumptions**: How these parameters calibrate the OFA Growth & Risk Simulator and IO Multiplier Engine.
4. **Identification Challenges & Uncertainty Analysis**: Sensitivity bounds, omitted variable caveats, and data limitations.
5. **Methodological References & Bibliography**: Full academic citations formatted in Canadian Journal of Agricultural Economics (CJAE) style.
"""
    else:  # Formal Government Submission / Policy Brief
        template_instructions = """
STRUCTURE AS A FORMAL OFA GOVERNMENT SUBMISSION / POLICY BRIEF:
1. **Title**: Policy Brief: [Descriptive Policy Title for Ontario Agriculture]
2. **Executive Summary**: 3-4 paragraph high-level synthesis with bolded core takeaways.
3. **Statement of the Issue & Ontario Agricultural Context**: Current market conditions, regulatory/trade environment, and provincial farm relevance.
4. **Empirical Literature Consensus & Evidence Table**: Comprehensive review of academic literature and grey literature. Include a structured Markdown comparison table.
5. **Economic & Farm Cash Flow Impact Analysis**: Empirical pass-through rates, gross margin implications, and multiplier effects across Ontario farm commodities.
6. **Policy Recommendations for Government**: 3-5 concrete, evidence-backed, actionable policy recommendations for OMAFA / AAFC.
7. **Academic & Grey Literature Bibliography**: Complete formal reference list (CJAE / APA style) linking all cited studies.
"""

    prompt = f"""You are the Lead Economic Policy Analyst at the Ontario Federation of Agriculture (OFA).
Draft a comprehensive, professional, highly rigorous policy document based on our curated research literature and empirical database.

PROJECT DOSSIER: {proj['name']}
PROJECT DESCRIPTION: {proj.get('description', '')}
TARGET AUDIENCE: {target_audience or BRIEF_TEMPLATES.get(template_type, {}).get('audience', 'Agricultural Policymakers')}
ADDITIONAL FOCUS / INSTRUCTIONS: {custom_focus or 'Ensure rigorous grounding in Ontario farm economic realities.'}

{template_instructions}

EMPIRICAL PARAMETERS AVAILABLE:
{params_text}

PRIOR LITERATURE SYNTHESIS:
{prior_synth_text}

CURATED RESEARCH PAPERS:
{papers_text}

WRITING GUIDELINES:
- Maintain an authoritative, evidence-based, professional tone suitable for senior agricultural leadership and government officials.
- Ground every claim strictly in the cited papers (cite as [1], [2], or Author, Year) and empirical parameters.
- Highlight specific Ontario agricultural commodities (grain/oilseeds, beef, hogs, supply-managed dairy/poultry, horticulture) where relevant.
- Do not use vague or generic policy statements; provide concrete economic mechanisms and policy designs.
"""

    response = model.generate_content(prompt)
    return response.text


def export_brief_to_docx(brief_markdown: str, title: str = "OFA Policy Brief", project_name: str = "") -> io.BytesIO:
    """Convert Markdown policy brief into a professionally formatted Word document (.docx)."""
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

    doc = docx.Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    green_primary = RGBColor(26, 86, 50)     # #1a5632 (OFA Green)
    dark_slate = RGBColor(30, 41, 59)        # #1e293b
    gold_accent = RGBColor(180, 83, 9)       # #b45309

    # Document Header Title
    header_p = doc.add_paragraph()
    header_run = header_p.add_run("ONTARIO FEDERATION OF AGRICULTURE")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(10)
    header_run.font.bold = True
    header_run.font.color.rgb = green_primary

    # Project Subheader
    if project_name:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"Research Dossier: {project_name}")
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(9)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Parse Markdown lines
    lines = brief_markdown.split("\n")
    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        rows_data = []
        for tl in t_lines:
            if re.match(r"^\s*\|?\s*[-:]+[-| :]+[-:]+\s*\|?\s*$", tl):
                continue  # Header separator line
            cols = [c.strip() for c in tl.strip().strip("|").split("|")]
            if cols:
                rows_data.append(cols)

        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        word_table = doc.add_table(rows=len(rows_data), cols=num_cols)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        word_table.style = "Table Grid"

        for r_idx, row in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row):
                if c_idx < num_cols:
                    cell = word_table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    # Header row styling
                    if r_idx == 0:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A5632"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(9.5)
                    else:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    for line in lines:
        stripped = line.strip()

        # Check for Table line
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        if not stripped:
            continue

        # Heading 1
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:].strip())
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = green_primary

        # Heading 2
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:].strip())
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = green_primary

        # Heading 3
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[4:].strip())
            run.font.name = "Arial"
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = gold_accent

        # Bullet List Item
        elif stripped.startswith("* ") or stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            # Parse bold runs inside bullets
            raw_text = stripped[2:].strip()
            parts = re.split(r"(\*\*.*?\*\*)", raw_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

        # Blockquote
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(stripped[2:].strip())
            r.font.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table(table_lines)

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes


def export_brief_to_html(brief_markdown: str, title: str = "OFA Policy Brief") -> str:
    """Convert Markdown policy brief into a styled, printable HTML document."""
    # Basic markdown-to-HTML conversion for headings, lists, tables, bold, italics
    html_body = brief_markdown
    
    # Convert headings
    html_body = re.sub(r"^### (.*?)$", r"<h3></h3>", html_body, flags=re.MULTILINE)
    html_body = re.sub(r"^## (.*?)$", r"<h2></h2>", html_body, flags=re.MULTILINE)
    html_body = re.sub(r"^# (.*?)$", r"<h1></h1>", html_body, flags=re.MULTILINE)
    
    # Convert bold and italic
    html_body = re.sub(r"\*\*(.*?)\*\*", r"<strong></strong>", html_body)
    html_body = re.sub(r"\*(.*?)\*", r"<em></em>", html_body)
    
    # Convert blockquotes
    html_body = re.sub(r"^> (.*?)$", r"<blockquote></blockquote>", html_body, flags=re.MULTILINE)
    
    # Convert bullets
    html_body = re.sub(r"^\* (.*?)$", r"<li></li>", html_body, flags=re.MULTILINE)
    html_body = re.sub(r"^- (.*?)$", r"<li></li>", html_body, flags=re.MULTILINE)
    
    # Convert paragraphs
    paragraphs = html_body.split("\n\n")
    formatted_p = []
    for p in paragraphs:
        p_str = p.strip()
        if p_str.startswith("<h") or p_str.startswith("<blockquote") or p_str.startswith("<li>"):
            formatted_p.append(p_str)
        elif "|" in p_str:
            formatted_p.append(p_str)  # Table text
        else:
            formatted_p.append(f"<p>{p_str}</p>")
    html_body = "\n".join(formatted_p)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
    @media print {{
        body {{ font-size: 11pt; }}
        .no-print {{ display: none; }}
    }}
    body {{
        font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, Roboto, Helvetica, Arial, sans-serif;
        line-height: 1.6;
        color: #1e293b;
        max-width: 850px;
        margin: 0 auto;
        padding: 2.5rem;
        background: #fff;
    }}
    .ofa-header {{
        border-bottom: 3px solid #1a5632;
        padding-bottom: 0.8rem;
        margin-bottom: 1.8rem;
        display: flex;
        justify-content: space-between;
        align-items: flex-end;
    }}
    .ofa-header h2 {{
        color: #1a5632;
        margin: 0;
        font-size: 1.3rem;
        font-weight: 700;
        letter-spacing: -0.3px;
    }}
    .ofa-header span {{
        color: #64748b;
        font-size: 0.85rem;
    }}
    h1 {{ color: #1a5632; font-size: 1.7rem; margin-top: 0; }}
    h2 {{ color: #1a5632; font-size: 1.25rem; border-bottom: 1px solid #e2e8f0; padding-bottom: 0.3rem; margin-top: 1.6rem; }}
    h3 {{ color: #b45309; font-size: 1.05rem; margin-top: 1.2rem; }}
    p {{ margin: 0.6rem 0; }}
    blockquote {{
        border-left: 4px solid #1a5632;
        background: #f8faf9;
        margin: 1rem 0;
        padding: 0.6rem 1rem;
        color: #334155;
        font-style: italic;
    }}
    li {{ margin-bottom: 0.35rem; }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 1.2rem 0;
        font-size: 0.9rem;
    }}
    th, td {{
        border: 1px solid #cbd5e1;
        padding: 8px 12px;
        text-align: left;
    }}
    th {{
        background: #1a5632;
        color: white;
        font-weight: 600;
    }}
    tr:nth-child(even) {{ background: #f8fafc; }}
</style>
</head>
<body>
<div class="ofa-header">
    <h2>ONTARIO FEDERATION OF AGRICULTURE</h2>
    <span>Policy & Economic Research Division</span>
</div>
{html_body}
</body>
</html>"""
