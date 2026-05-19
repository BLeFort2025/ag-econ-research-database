"""Build the updated Action Plan Word doc with Round 2 expert corrections."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = r"C:\Users\ben.lefort\OneDrive - Ontario Federation of Agriculture\Desktop\Ben Desktop Files\Economic Analyst Position\Economic papers\Ag Economic Research Database\Local Food"
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1B, 0x3A, 0x2D)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def add_note(text, label="NOTE"):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p.add_run(text)

def bold_para(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

# ═══════════════════════════════════════════════════════════════════════════
#  TITLE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Action Plan: Building Ontario Local Food Multipliers")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x2D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("What We Have, What We Need, and the Path to Publication")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("March 2026 | Incorporating Two Rounds of Expert Peer Review")
run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 1: WHAT WE HAVE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 1: What We Have Built', level=1)

doc.add_paragraph(
    'The Local Food Multiplier Engine (local_food_multiplier.py) implements a 5-phase '
    'Analysis-by-Parts pipeline using corrected methodology from two rounds of expert review:'
)

doc.add_paragraph(
    '1. Custom DTC farm production function (how local farms spend each dollar)\n'
    '2. Purchaser Price to Basic Price margin deflation (StatCan SUT architecture)\n'
    '3. Regional Purchase Coefficient (RPC) application (local vs. imported)\n'
    '4. Gross impact via sector-specific Simple multipliers (no double-counting)\n'
    '5. Swenson Net-Impact counterfactual (displaced conventional grocery activity)'
)

doc.add_heading('Prototype Results (Mock Data)', level=2)
doc.add_paragraph('Per $1,000,000 of consumer food spending redirected from conventional to local:')

add_table(
    ["Metric", "Naive StatCan Baseline", "Gross Local Food", "Net Local Food"],
    [
        ["Total Output", "$1.82M", "$1.35M", "$531K"],
        ["GDP Contribution", "$450K", "$830K (+84%)", "$451K"],
        ["Jobs Supported", "8.5", "15.1 (+78%)", "7.1"],
        ["Output Multiplier", "1.82x", "1.35x", "0.53x"],
    ]
)

doc.add_paragraph(
    '\nExpert validation: "The AI perfectly anticipated the expected outcome. It is completely '
    'normal and mathematically correct for the Gross Output Multiplier to drop (1.82x to 1.35x) '
    'because local farms buy fewer heavy manufactured inputs. The fact that the engine produced '
    'this inverse relationship proves its logic is wired structurally, not just scaling numbers up."'
)

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 2: EXPERT CORRECTIONS APPLIED
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 2: Expert Corrections Applied (Two Rounds)', level=1)

doc.add_heading('Round 1: Structural I-O Corrections (Implemented in Engine)', level=2)

bold_para('Correction 1: No Scalar Adjustment to Multipliers')
doc.add_paragraph(
    'WRONG: Adjusted_Multiplier = StatCan_Multiplier x AdjustmentFactor\n'
    'RIGHT: Bottom-up vector shock using Analysis-by-Parts. The engine distributes each '
    'dollar through specific IOIC sectors, applies sector-specific Simple multipliers, and '
    'sums the individual impacts. STATUS: IMPLEMENTED.'
)

bold_para('Correction 2: Purchaser to Basic Price Margin Deflation')
doc.add_paragraph(
    'WRONG: Shocking Paper Mfg with the full packaging purchase price.\n'
    'RIGHT: Strip retail, wholesale, and transport margins before applying manufacturing '
    'multiplier. A $100 packaging purchase = $20 retail + $10 wholesale + $5 transport + '
    '$65 paper mill. STATUS: IMPLEMENTED.'
)

bold_para('Correction 3: Simple Multipliers Only')
doc.add_paragraph(
    'WRONG: Using "Total" multipliers then separately calculating Induced effects.\n'
    'RIGHT: Use SIMPLE multipliers for supply chain, manually calculate Induced effects '
    'via Household Consumption vector with 25% tax/savings leakage. STATUS: IMPLEMENTED.'
)

doc.add_heading('Round 2: Advanced Data Wrangling Corrections (To Apply in Phase 1-2)', level=2)

bold_para('Correction 4: The "Non-Cash" Trap in OMAFRA Enterprise Budgets')
doc.add_paragraph(
    'ISSUE: OMAFRA enterprise budgets include non-cash entries like Depreciation on '
    'tractors/buildings and Imputed Operator Labor (valuing farmer time at a theoretical '
    'hourly rate even if no cash was drawn).\n\n'
    'FIX: I-O models are driven strictly by actual cash flows.\n'
    '  - STRIP Depreciation entirely before converting to dollar fractions\n'
    '  - RECLASSIFY Imputed Operator Labor as Proprietor Income (routes to Household '
    'sector for induced effects) based on actual cash draws or net farm income\n\n'
    'STATUS: TO IMPLEMENT when ingesting OMAFRA data.'
)

bold_para('Correction 5: International vs. Interprovincial RPCs')
doc.add_paragraph(
    'ISSUE: StatCan Table 36-10-0612-01 (Interprovincial Trade Flows) only shows trade '
    'between Canadian provinces. Ontario imports massive amounts of conventional food and '
    'ag inputs from the USA and Mexico. Using only interprovincial data vastly underestimates '
    'import leakage.\n\n'
    'FIX: Must pull International Import vectors from the main Supply and Use Tables.\n\n'
    'FORMULA:\n'
    '  True RPC = Ontario Production consumed in Ontario /\n'
    '             (Ontario Production + Interprovincial Imports + International Imports)\n\n'
    'STATUS: TO IMPLEMENT when building real RPCs.'
)

bold_para('Correction 6: StatCan L-Level vs. W-Level Aggregation')
doc.add_paragraph(
    'ISSUE: StatCan publishes Table 36-10-0595-01 at the Summary (S) or Link (L) level. '
    'You will likely find a generic multiplier for BS1110 (Crop Production) or, at best, '
    'BS1112 (Vegetable and melon farming). These may be too aggregated.\n\n'
    'FIX: Check early in Phase 1 if publicly available L-level IOIC codes are granular '
    'enough for your mapped inputs. If they lump greenhouse vegetables with field cash crops, '
    'you may need to submit a custom data request to StatCan Industry Accounts Division '
    'for suppressed Worksheet (W) level tables. Budget: potentially $500-2,000 for a custom tab.\n\n'
    'STATUS: CHECK in Phase 1, Week 1.'
)

bold_para('Correction 7: The "Leaky Counterfactual" (Swenson Refinement)')
doc.add_paragraph(
    'ISSUE: The Swenson counterfactual currently assumes that conventional grocery imports '
    '80% of farm-gate product from outside Ontario. But Loblaws, Metro, etc. actually procure '
    'significant Ontario food (greenhouse tomatoes, dairy, poultry under supply management).\n\n'
    'FIX: Use Ontario agricultural self-sufficiency ratios to set the Farm-Gate RPC for the '
    'negative shock. Acknowledge that buying local may displace some conventional Ontario farm '
    'jobs, not just imported product. This makes the model more conservative but more defensible.\n\n'
    'EFFECT: Net impact will be somewhat lower than current mock results, but still positive '
    'and now peer-review-proof.\n\n'
    'STATUS: TO IMPLEMENT when calibrating real RPCs.'
)

bold_para('Correction 8: OFA Survey Design Pro-Tip')
doc.add_paragraph(
    'ISSUE: Asking farmers for absolute dollar amounts (e.g., "How much did you spend on fuel?") '
    'will crash response rates due to financial privacy concerns.\n\n'
    'FIX: Ask for expenditure PERCENTAGES instead:\n'
    '  "Think about your total farm revenue last year. Roughly what percentage went to:\n'
    '   Hired Labor? Paying yourself? Seed/Fertilizer? Fuel? Mortgage/Taxes?"\n\n'
    'This directly gives you the $1.00 vector fractions needed for the Leontief A-matrix. '
    'It is less invasive, higher response rate, and scales perfectly into the Python engine.\n\n'
    'STATUS: APPLY when designing survey instrument.'
)

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 3: WHAT WE STILL NEED
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 3: What We Still Need', level=1)

doc.add_heading('A. Statistics Canada Data (Replace Mock Multipliers)', level=2)

add_table(
    ["Priority", "Data Item", "StatCan Table", "Cost", "Purpose"],
    [
        ["CRITICAL", "Provincial Simple & Total Multipliers (Ontario)", "36-10-0595-01", "Free", "Replace ALL mock multipliers; derive Implied Induced multiplier"],
        ["CRITICAL", "Symmetric Input-Output Tables (Ontario)", "36-10-0478-01", "Free", "Source real margin profiles (Purchaser to Basic Price conversion)"],
        ["HIGH", "Provincial Supply & Use Tables", "36-10-0580-01", "Free", "Detailed industry purchase vectors for margin stripping"],
        ["HIGH", "Interprovincial Trade Flows", "36-10-0612-01", "Free", "Interprovincial component of RPC calculation"],
        ["HIGH", "International Import Vectors (from SUTs)", "Embedded in 36-10-0580-01", "Free", "International component of RPC (Correction #5)"],
        ["MEDIUM", "Household Final Consumption vector", "36-10-0107-01", "Free", "Properly model induced effects"],
        ["CONDITIONAL", "W-Level (suppressed) multipliers", "Custom StatCan request", "$500-2,000", "Only if L-level aggregation is too coarse (Correction #6)"],
    ]
)

doc.add_heading('B. Ontario Farm Budget Data (Replace Mock Production Function)', level=2)

add_table(
    ["Priority", "Data Item", "Source", "Method", "Correction Notes"],
    [
        ["CRITICAL", "DTC Farm Cost-of-Production", "OMAFRA Enterprise Budgets", "Download from OMAFRA website", "Strip depreciation & reclassify imputed labor (Correction #4)"],
        ["CRITICAL", "DTC Farmer % Expenditure Survey", "OFA Member Engagement", "Design survey (ask % not $)", "Use percentage questions per Correction #8"],
        ["HIGH", "Ontario ag self-sufficiency ratios", "OMAFRA / StatCan", "Commodity-level production vs. consumption", "For leaky counterfactual (Correction #7)"],
        ["HIGH", "CSA/Farm Stand Financial Data", "OFA + Ecological Farmers", "Request anonymized data", "Strip non-cash items per Correction #4"],
        ["MEDIUM", "Census of Agriculture Microdata", "StatCan (special tab)", "$$ (apply for access)", "Granular farm-level cost data by marketing channel"],
    ]
)

doc.add_heading('C. Literature to Acquire', level=2)

add_table(
    ["Priority", "Reference", "Where to Find It", "Purpose"],
    [
        ["CRITICAL", "Jablonski, Schmit & Mansury (2016)", "Google Scholar / JAFIO", "Custom DTC production function methodology"],
        ["CRITICAL", "USDA Local Food Economics Toolkit", "USDA AMS website (free)", "Excel templates for Analysis-by-Parts"],
        ["HIGH", "Swenson, D. - Iowa State reports", "Iowa State Extension", "Net impact / import substitution methodology"],
        ["HIGH", "Harry Cummings - Ontario Farmers Market Studies", "U of Guelph / FMO", "Only Canadian precedent for farmers market I-O"],
        ["HIGH", "Pacific Analytics - Feed BC (2021)", "BC Government", "Canadian I-O model for local food procurement"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 4: UPDATED 16-WEEK ROADMAP
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 4: Updated 16-Week Roadmap', level=1)

doc.add_heading('Phase 1: Data Collection & Ingestion (Weeks 1-4)', level=2)

bold_para('Week 1: StatCan Data Pipeline')
doc.add_paragraph(
    '1. Download StatCan Table 36-10-0595-01 (Ontario multipliers, CSV)\n'
    '2. Write Python ingestion script to parse and pivot into clean DataFrame\n'
    '3. Filter Ontario, isolate Simple + Total multipliers for Output, GDP, Jobs\n'
    '4. Calculate Implied Induced Multiplier = Total - Simple for each sector\n'
    '5. CHECK: Are L-level IOIC codes granular enough? (Correction #6)\n'
    '   If not: submit StatCan custom data request for W-level tables'
)

bold_para('Week 2: Margin Profiles & International RPCs')
doc.add_paragraph(
    '1. Download StatCan SUTs (36-10-0478-01, 36-10-0580-01)\n'
    '2. Extract margin profiles: how much of each product purchase is\n'
    '   retail/wholesale/transport/basic price\n'
    '3. Extract BOTH interprovincial AND international import vectors\n'
    '4. Calculate TRUE Ontario RPCs per Correction #5:\n'
    '   RPC = ON production in ON / (ON production + interp. imports + intl. imports)\n'
    '5. Source Ontario agricultural self-sufficiency ratios for Correction #7'
)

bold_para('Week 3: OMAFRA Enterprise Budgets')
doc.add_paragraph(
    '1. Download OMAFRA enterprise budgets for: fresh market vegetables,\n'
    '   berries, greenhouse crops, mixed farms\n'
    '2. STRIP non-cash items per Correction #4:\n'
    '   - Remove Depreciation entirely\n'
    '   - Reclassify Imputed Operator Labor as Proprietor Income\n'
    '3. Convert remaining cash costs to $1.00 vector fractions\n'
    '4. Build 2-3 preliminary DTC production functions'
)

bold_para('Week 4: OFA Member Survey')
doc.add_paragraph(
    '1. Design short (8-10 question) expenditure PERCENTAGE survey (Correction #8)\n'
    '2. Sample questions:\n'
    '   "What % of total revenue went to: Hired Labor? Your own draw?\n'
    '    Seed/Fertilizer? Fuel? Marketing? Taxes/Mortgage?"\n'
    '3. Deploy to OFA members who sell direct-to-consumer\n'
    '4. Target: 50-100 responses across farmers market, CSA, farm stand channels\n'
    '5. Begin collecting responses (allow 3-4 weeks for returns)'
)

doc.add_heading('Phase 2: Engine Calibration (Weeks 5-8)', level=2)
doc.add_paragraph(
    '1. Replace mock multipliers with real StatCan Simple multipliers\n'
    '2. Replace mock margin profiles with real SUT-sourced profiles\n'
    '3. Replace proxy RPCs with calculated TRUE Ontario RPCs\n'
    '4. Build 2-3 real DTC production functions from OMAFRA + survey data\n'
    '5. Refine Swenson counterfactual with real self-sufficiency ratios (Correction #7)\n'
    '6. Run sensitivity analysis on RPCs, margins, and budget assumptions\n'
    '7. Document all data sources, transformations, and assumptions'
)

doc.add_heading('Phase 3: Validation & Peer Review (Weeks 9-12)', level=2)
doc.add_paragraph(
    '1. Compare results against US benchmarks:\n'
    '   - Jablonski: employment mult ~2.5-3.5x for DTC farms\n'
    '   - Oklahoma IMPLAN: ~2.4x total output for farmers markets\n'
    '2. Compare against Cummings Ontario farmers market studies (1.5-3.0x range)\n'
    '3. Compare against Feed BC report methodology and results\n'
    '4. Write technical methodology report for OFA executive\n'
    '5. Submit to external ag economist (University of Guelph) for peer review\n'
    '6. Prepare plain-language advocacy brief for OFA Board'
)

doc.add_heading('Phase 4: Dashboard Integration (Weeks 13-16)', level=2)
doc.add_paragraph(
    '1. Add "Local Food Impact" module to the Agri-Food Economic Dashboard\n'
    '2. Allow users to select farm type (vegetable, berry, CSA, mixed)\n'
    '3. Allow users to input local food spending amount\n'
    '4. Display Gross vs. Net impact comparison with interactive charts\n'
    '5. Generate consulting-grade PDF report for municipal delegations\n'
    '6. Build "Foodland Ontario Impact Calculator" for public-facing advocacy'
)

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 5: MONDAY STARTING POINT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Section 5: Monday Starting Point - Phase 1 Data Ingestion", level=1)

doc.add_paragraph(
    'The immediate next step is to get real Statistics Canada data flowing into the engine. '
    'On Monday, the first task is:'
)

bold_para('Task 1: Download & Ingest StatCan Table 36-10-0595-01')
doc.add_paragraph(
    '1. Download the CSV from StatCan open data portal\n'
    '2. Write Python ingestion script to:\n'
    '   a. Filter for Geography = "Ontario"\n'
    '   b. Pivot into clean DataFrame with IOIC codes as index\n'
    '   c. Columns: Simple Output, Simple GDP, Simple Jobs,\n'
    '               Total Output, Total GDP, Total Jobs\n'
    '   d. Calculate: Implied Induced = Total - Simple for each sector\n'
    '3. Verify key sectors are available:\n'
    '   Crop/Animal Production, Retail Trade, Wholesale Trade,\n'
    '   Truck Transport, Paper Mfg, Petroleum, Professional Services\n'
    '4. Plug real multipliers into local_food_multiplier.py'
)

bold_para('Task 2: Check IOIC Aggregation Level')
doc.add_paragraph(
    '1. Review the industry classifications available at the L-level\n'
    '2. Determine if "Vegetable and melon farming" (BS1112) exists separately\n'
    '3. If too aggregated: draft custom data request letter for StatCan\n'
    '4. Document any gaps or aggregation concerns'
)

add_note(
    'The PDF download script should be complete or near-complete by Monday. Check its '
    'status and run the text extraction pipeline on any newly downloaded papers to enrich '
    'the database for further literature research.',
    'REMINDER'
)

# ── Save ─────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "Action Plan - What We Still Need.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
