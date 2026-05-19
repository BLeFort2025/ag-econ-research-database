"""Convert the local food multiplier research briefing to a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = r"C:\Users\ben.lefort\OneDrive - Ontario Federation of Agriculture\Desktop\Ben Desktop Files\Economic Analyst Position\Economic papers\Ag Economic Research Database\Local Food"
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x2D)  # Dark green

# ── Helper functions ─────────────────────────────────────────────────────
def add_bold_para(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_note(text, label="NOTE"):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p.add_run(text)
    return p

# ── Title Page ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Adjusting Statistics Canada I-O Multipliers\nfor Local Food Economic Impact")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x2D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Research Briefing — Deep Database Sweep")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("March 2026").font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Database: 54,071 papers searched across 15 research themes\n3,903 papers match at least one theme • 353 match 2+ themes")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── Executive Summary ────────────────────────────────────────────────────
doc.add_heading('Executive Summary', level=1)

doc.add_paragraph(
    "Your database contains strong building blocks for a made-in-Canada local food multiplier methodology - "
    'but the directly applicable literature is thin (only 4 papers sit at the intersection of "local food" '
    'AND "I-O methodology"). The critical gap is that the US literature (Jablonski, Swenson, Thilmany) is '
    "mostly in journals your database covers (AJAE, AEPP, Food Policy), but those specific landmark papers "
    "may not have been captured by your harvesting keywords."
)

doc.add_paragraph(
    "The good news: you have 140 I-O papers, 115 multiplier papers, 32 IMPLAN papers, 48 SAM papers, "
    "67 local food papers, and 32 farmers market papers — more than enough to build the methodology "
    "from component parts."
)

# ── Section 1: Highest-Value Papers ─────────────────────────────────────
doc.add_heading('Section 1: Highest-Value Papers in Your Database', level=1)

doc.add_heading('Tier A — Directly On-Point', level=2)
add_table(
    ["#", "Year", "Cited", "Title", "DOI", "Notes"],
    [
        ["1", "2008", "0", "Economic Impacts of Direct Produce Marketing: Oklahoma Farmers' Markets", "10.22004/ag.econ.6785", "IMPLAN model applied to farmers markets"],
        ["2", "2023", "9", "Economic impact assessment of public incentives: farm-to-school", "10.1016/j.foodpol.2023.102545", "NY 30% Initiative — detailed school purchasing data"],
        ["3", "2015", "56", "Economic impact of farmers' markets and locally grown campaign", "10.1016/j.foodpol.2015.05.001", "Direct relevance to Foodland Ontario"],
        ["4", "2014", "78", "Economic impact of direct marketing: Vietnam vegetable chains", "10.1016/j.foodpol.2014.04.001", "Direct marketing economic impact framework"],
    ]
)

doc.add_heading('Tier B — Critical Supporting Papers', level=2)
add_table(
    ["#", "Year", "Cited", "Title", "DOI", "Why It Matters"],
    [
        ["5", "2003", "0", "Regional Economic Impacts of Florida Agriculture", "10.22004/ag.econ.15702", "Full IMPLAN + I-O methodology template"],
        ["6", "2003", "0", "IMPLAN Understates Agricultural I-O Multipliers: Colorado", "10.22004/ag.econ.14669", "Documents multiplier underestimation — supports your case"],
        ["7", "2020", "10", "Local Foods and Local Economic Performance", "10.1002/aepp.13015", "DTC + local food economic performance — AEPP"],
        ["8", "2017", "38", "Retail Intermediation and Local Foods", "10.1093/ajae/aaw115", "AJAE: margin reallocation in local food retail"],
        ["9", "2020", "20", "Local Foods Go Downstream: Spatial Factors", "10.1002/aepp.13046", "Local food downstream manufacturing effects"],
        ["10", "2021", "20", "Novel methods for local food systems' impacts", "10.1111/1467-8489.12456", "Methodological innovation for local food impacts"],
        ["11", "2012", "140", "Energy efficiency of local food systems", "10.1016/j.foodpol.2012.07.006", "Highly cited local food system analysis"],
        ["12", "2015", "53", "How Local Is Local? Canadian Food Labeling Policy", "10.1111/cjag.12062", "CJAE — defines 'local' for Canadian policy"],
        ["13", "2024", "5", "Price changes: local food vs. mainstream grocery in Canada", "10.1016/j.foodpol.2024.102773", "Canadian local food pricing — recent"],
        ["14", "2013", "78", "Do farm operators benefit from DTC marketing?", "10.1111/agec.12042", "DTC farm economics"],
    ]
)

doc.add_heading('Tier C — I-O Methodology Building Blocks', level=2)
add_table(
    ["Category", "Count", "Key Application"],
    [
        ["Input-Output methodology", "140", "Canadian, US, OECD applications across ag sectors"],
        ["IMPLAN applications", "32", "Florida, Colorado, Oklahoma, Idaho — portable methodologies"],
        ["SAM (Social Accounting Matrix)", "48", "Rural CAP analysis, poverty multiplier decomposition"],
        ["Leontief framework", "29", "Core methodology including Canadian applications"],
        ["Supply and Use Tables", "7", "Directly relevant to StatCan SUT adaptation"],
        ["Employment multipliers", "6", "Farm labor multiplier calculations"],
        ["Production functions", "500", "Massive resource for custom local farm production functions"],
        ["Enterprise budgets", "19", "Direct farm cost structure data"],
        ["Marketing margins", "70", "Retail/wholesale margin analysis — key for margin reallocation"],
        ["Import substitution", "28", "Relevant to Swenson 'net impact' approach"],
    ]
)

# ── Section 2: Literature Gaps ──────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Section 2: Critical Literature Gaps (Papers to Acquire)', level=1)

doc.add_paragraph(
    "These landmark papers are NOT in your database but are essential for the methodology:"
)

add_table(
    ["#", "Authors", "Year", "Title", "Source", "Why Critical"],
    [
        ["1", "Jablonski, Schmit & Mansury", "2016", "Assessing Economic Impacts of Local Food System Producers by Scale", "J. of Ag & Food Industrial Org.", "THE landmark paper — built custom 'Small Direct Agriculture' production function for IMPLAN"],
        ["2", "Thilmany, Jablonski, Swenson et al.", "2016", "USDA Economics of Local Food Systems Toolkit", "USDA AMS", "Standardized the Analysis-by-Parts method"],
        ["3", "Swenson, David", "Various", "Iowa State reports on local food import substitution", "Iowa State Extension", "Pioneer of 'net impact' methodology"],
        ["4", "Harry Cummings & Assoc.", "1999–2015", "Ontario Farmers' Market Economic Impact Studies", "U of Guelph (commissioned)", "Only Canadian precedent for rigorous farmers market I-O"],
        ["5", "Pacific Analytics", "2021", "Feed BC Economic Impacts Report", "BC Gov't commissioned", "Canadian provincial I-O model for local food procurement"],
    ]
)

add_note(
    "The Jablonski et al. (2016) paper is the single most important reference. It provides the exact "
    "methodology for building a custom production function to replace StatCan's generic agricultural multiplier. "
    "Search Google Scholar or contact the authors directly.",
    "CRITICAL"
)

# ── Section 3: Actionable Steps ─────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Section 3: Actionable Steps — Building Your Local Food Multiplier', level=1)

# Step 1
doc.add_heading('Step 1: Define "Local" for Ontario (Week 1)', level=2)
doc.add_paragraph('Database resource: Paper #12 (How Local is Local? CJAE 2015, 53 citations)')
doc.add_paragraph(
    '• Establish your geographic boundary (provincial? 100km? county?)\n'
    '• The CJAE paper specifically examines Canadian consumer perceptions of "local food" labeling\n'
    '• Align with Foodland Ontario\'s existing definition for policy credibility'
)

# Step 2
doc.add_heading('Step 2: Build the Local Farm Production Function (Weeks 1–3)', level=2)
doc.add_paragraph('Database resources: 500 production function papers, 19 enterprise budget papers, 70 cost structure papers')
doc.add_paragraph(
    'This is the core technical step. You need to answer: "When a local Ontario farm earns $1, where does each cent go?"'
)

p = add_bold_para('Data sources to build this:')
doc.add_paragraph(
    '1. Ontario Ministry of Agriculture enterprise budgets — published for major crops\n'
    '2. OMAFRA cost-of-production surveys — primary data on Ontario farm expenses\n'
    '3. Census of Agriculture microdata (if accessible via StatCan)\n'
    '4. Direct surveys of OFA member farms selling locally (farmers markets, CSAs, farm stands)'
)

add_table(
    ["Category", "Conventional Farm", "Local/DTC Farm", "Source"],
    [
        ["Labor (hired + family)", "~15–20%", "~30–40%", "Enterprise budgets; Paper #14"],
        ["Proprietor income", "~10%", "~25–35%", "Your survey data"],
        ["Seed/fertilizer/chemicals", "~25–30%", "~10–15%", "Enterprise budgets"],
        ["Packaging/marketing", "~2%", "~8–12%", "Survey"],
        ["Equipment/machinery", "~15–20%", "~5–10%", "Enterprise budgets"],
        ["Property taxes/land", "~5–8%", "~5–8%", "MPAC data"],
        ["Transportation", "~5–8%", "~2–3%", "Survey"],
        ["Wholesale/retail margins", "~25–35% (leaks)", "~0–5% (farmer captures)", "Papers #8, #11"],
    ]
)

# Step 3
doc.add_heading('Step 3: Map to StatCan IOIC Codes (Week 3)', level=2)
doc.add_paragraph('Database resources: 140 I-O papers, 7 supply & use table papers')
doc.add_paragraph('For each local farm expenditure line, map to the StatCan Input-Output Industry Classification (IOIC) code:')

add_table(
    ["Farm Expenditure", "StatCan IOIC Code", "Multiplier Table"],
    [
        ["Local labor", "Household sector", "StatCan household consumption multipliers"],
        ["Seed/inputs", "Crop production (BS111)", "Table 36-10-0594-01"],
        ["Packaging", "Paper product mfg (BS3222)", "Standard multiplier"],
        ["Fuel/energy", "Petroleum products (BS3241)", "Standard multiplier"],
        ["Equipment repair", "Machinery mfg (BS3331)", "Standard multiplier"],
        ["Property taxes", "Government sector", "Provincial/municipal redistribution"],
    ]
)

# Step 4
doc.add_heading('Step 4: Run Analysis-by-Parts (Week 4)', level=2)
doc.add_paragraph('Database resources: Paper #5 (Florida), Paper #6 (IMPLAN understates multipliers)')
doc.add_paragraph(
    'Instead of applying StatCan\'s generic "Crop production" multiplier to $X of local food sales, you:\n\n'
    '1. Decompose the $X into each expenditure line from Step 2\n'
    '2. Apply the RPC (Regional Purchase Coefficient) — what % is bought in-province?\n'
    '3. Multiply each locally-purchased input by its specific StatCan multiplier\n'
    '4. Sum the individual impacts = your custom Direct + Indirect effect\n'
    '5. Calculate Induced effects separately using household consumption multipliers'
)
add_bold_para('Formula:')
doc.add_paragraph(
    'Custom Local Food Multiplier = Σ(ExpenditureShare_i × RPC_i × StatCanMultiplier_i) + InducedEffect'
)

# Step 5
doc.add_heading('Step 5: Calculate Margin Reallocation (Week 4)', level=2)
doc.add_paragraph('Database resources: 70 marketing margin papers, 15 food retail margin papers, Paper #8')
doc.add_paragraph('This is where local food creates the biggest multiplier difference:')

add_table(
    ["Supply Chain", "Farm Gate", "Wholesale/Distribution", "Retail", "Processing"],
    [
        ["Conventional ($1)", "~$0.15", "~$0.25 (often out-of-province)", "~$0.35 (corporate HQ often out-of-province)", "~$0.25"],
        ["Local DTC ($1)", "~$0.70–0.85 (captures retail margin)", "~$0.00", "~$0.00", "~$0.15–0.30 (local)"],
    ]
)

doc.add_paragraph(
    'The local farmer capturing the retail margin means that money stays in the regional economy instead '
    'of leaking to corporate supply chains — this is where the enhanced multiplier comes from.'
)

# Step 6
doc.add_heading('Step 6: Apply the Swenson Net Impact Correction (Week 5)', level=2)
doc.add_paragraph('Database resources: 28 import substitution papers, 17 net impact papers, 121 counterfactual papers')
doc.add_paragraph('This is mandatory for academic credibility. Run two scenarios:')

add_table(
    ["Scenario", "Calculation"],
    [
        ["Positive shock", "Local food purchases × your custom multiplier from Steps 4–5"],
        ["Negative shock", "Same dollar amount × conventional grocery retail StatCan multiplier"],
        ["Net Impact", "Positive − Negative = True economic gain from buying local"],
    ]
)

doc.add_paragraph(
    'The literature (Swenson) consistently shows that net impacts are positive but smaller than gross impacts — '
    'typically 40–60% of the gross figure.'
)

# Step 7
doc.add_heading('Step 7: Validate Against Existing Studies (Week 6)', level=2)
doc.add_paragraph('Database resources: Paper #1 (Oklahoma IMPLAN: $3.3M → $7.8M total impact, ~2.4× multiplier)')
doc.add_paragraph(
    'Your results should produce:\n'
    '• Output multiplier: 1.4–2.0× (vs. StatCan generic crop: ~1.8–2.2×)\n'
    '• Employment multiplier: 2.0–3.5× (significantly higher due to labor intensity)\n'
    '• Income multiplier: 1.5–2.5× (higher due to margin capture)\n\n'
    'If your employment and income multipliers are higher than StatCan\'s generic agricultural multipliers '
    'but total output is similar or lower, you\'ve replicated the core finding from Jablonski et al. (2016).'
)

# ── Section 4: Dashboard Integration ────────────────────────────────────
doc.add_page_break()
doc.add_heading('Section 4: Integration with Your Dashboard', level=1)

doc.add_paragraph(
    'Based on the dashboard\'s existing I-O engine architecture (the $338B baseline), you could add a '
    '"Local Food Premium" adjustment layer:'
)

add_bold_para('Adjusted_Multiplier = StatCan_Base_Multiplier × LocalFoodAdjustmentFactor')

doc.add_paragraph('Where the adjustment factor accounts for:')
doc.add_paragraph(
    '1. Labour intensity premium: +40–80% on employment multiplier\n'
    '2. Margin capture premium: +20–50% on income multiplier\n'
    '3. Leakage reduction: −10–30% on import content\n'
    '4. Net impact discount: −40–60% (Swenson correction for displaced conventional spending)'
)

# ── Section 5: Recommended Next Actions ──────────────────────────────────
doc.add_heading('Section 5: Recommended Next Actions', level=1)

doc.add_paragraph(
    '1. Immediate: Download and read the Oklahoma IMPLAN farmers market paper (10.22004/ag.econ.6785) — '
    'it\'s in AgEcon Search and should download via your running script\n'
    '2. This week: Acquire the Jablonski et al. (2016) paper from JAFIO — this is the methodological keystone\n'
    '3. This week: Read the CJAE "How Local is Local?" paper (2015) for Canadian definitional context\n'
    '4. Next 2 weeks: Build the Ontario local farm production function using OMAFRA enterprise budgets + OFA member survey data\n'
    '5. Month 2: Run the Analysis-by-Parts calculation using StatCan Table 36-10-0594-01 multipliers\n'
    '6. Month 2: Apply net impact correction and validate against US benchmarks'
)

add_note(
    "The USDA's Local Food Economics Toolkit (free, online) provides Excel spreadsheets that implement "
    "Analysis-by-Parts. You can use these as templates and swap in StatCan multiplier values for the IMPLAN values. "
    "This is by far the fastest path to a working prototype.",
    "TIP"
)

# ── Appendix ─────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Appendix: Full Search Statistics', level=1)

search_stats = [
    ("Canadian Agriculture Impact", 500), ("Food Security Economic", 500),
    ("Economic Impact Agriculture", 500), ("Production Function", 500),
    ("Opportunity Cost", 247), ("Value Chain Food", 243),
    ("Agri-food", 206), ("Ontario Food", 192),
    ("Input-Output Food", 140), ("Counterfactual", 121),
    ("Multiplier Agriculture", 115), ("Displacement Effect", 105),
    ("Agrifood", 89), ("Rural Economy", 77),
    ("Regional Economic Impact", 75), ("Cost Structure Farm", 70),
    ("Marketing Margin", 70), ("Local Food Economics", 67),
    ("Rural Economic", 55), ("Social Accounting Matrix", 48),
    ("Direct Marketing Farm", 32), ("IMPLAN", 32),
    ("Food Distribution", 30), ("Leontief", 29),
    ("Import Substitution", 28), ("Enterprise Budget", 19),
    ("Community Economic", 19), ("Farmers Markets", 17),
    ("Net Impact", 17), ("Farmers Market", 15),
    ("Food Retail Margin", 15), ("Local Food Systems", 9),
    ("Buy Local", 9), ("Locally Grown", 9),
    ("Public Procurement Food", 9), ("Food Miles", 7),
    ("Supply Use Table", 7), ("Employment Multiplier", 6),
    ("Direct to Consumer", 5), ("Short Supply Chain", 5),
    ("Community Food", 5), ("Farm to School", 4),
    ("Food Hub", 3), ("Direct-to-Consumer", 3),
    ("Farm to Table", 3), ("Economic Multiplier", 2),
    ("Local Procurement Food", 2), ("Foodland Ontario", 1),
]

add_table(
    ["Search Theme", "Papers Found"],
    [[s[0], f"{s[1]:,}"] for s in search_stats]
)

# ── Save ─────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "Local Food Multiplier Research Briefing.docx")
doc.save(out_path)
print(f"Saved to: {out_path}")
